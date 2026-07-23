"""Create ATS-friendly CV data and downloadable PDF/DOCX exports."""
import io
import json
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


def _clean_json(text: str) -> dict[str, Any]:
    text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.I)
    start, end = text.find("{"), text.rfind("}")
    if start < 0 or end < 0:
        raise ValueError("The model did not return a usable CV structure. Please try again.")
    return json.loads(text[start:end + 1])


def generate_resume(rag, profile: dict[str, Any], role: str, sources: list[dict], language: str = "English") -> dict[str, Any]:
    context = "\n\n".join(f"[{i}] {x['text']}" for i, x in enumerate(sources, 1))
    instruction = f"""Create an ATS-friendly, truthful CV for the target role: {role}.
CV language: {language}. Use the supplied candidate data only; never invent employers, dates, degrees, metrics, achievements, certificates, or contact data. Improve wording and organize relevant details for the role using the retrieved career context. Write strong concise accomplishment bullets when supported by the candidate input. If an area is missing, omit it.
Return ONLY valid JSON with exactly these keys:
{{"name":"", "headline":"", "contact":[""], "summary":"", "skills":[""], "experience":[{{"title":"","company":"","dates":"","bullets":[""]}}], "education":[{{"school":"","degree":"","dates":""}}], "projects":[{{"name":"","link":"","bullets":[""]}}], "links":[{{"label":"","url":""}}]}}

Candidate data:\n{json.dumps(profile, ensure_ascii=False)}\n\nRetrieved career context:\n{context}"""
    if not rag.OPENROUTER_API_KEY:
        raise RuntimeError("Missing OPENROUTER_API_KEY. Add it to Streamlit secrets or your environment.")
    import requests
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={"Authorization": f"Bearer {rag.OPENROUTER_API_KEY}", "Content-Type": "application/json"},
        json={"model": rag.OPENROUTER_MODEL, "messages": [{"role": "system", "content": "You are a precise professional resume writer. Output valid JSON only."}, {"role": "user", "content": instruction}], "temperature": 0.15},
        timeout=60,
    )
    response.raise_for_status()
    return _clean_json(response.json()["choices"][0]["message"]["content"])


def _esc(value: Any) -> str:
    from html import escape
    return escape(str(value or ""))


def resume_html(cv: dict[str, Any], accent: str = "#4f46e5") -> str:
    def section(title, content):
        return f"<h2>{title}</h2>{content}" if content else ""
    experience = "".join(f"<div class='item'><b>{_esc(x.get('title'))}</b> <span>· {_esc(x.get('company'))}</span><em>{_esc(x.get('dates'))}</em><ul>{''.join(f'<li>{_esc(b)}</li>' for b in x.get('bullets', []))}</ul></div>" for x in cv.get("experience", []))
    education = "".join(f"<div class='item'><b>{_esc(x.get('school'))}</b><em>{_esc(x.get('dates'))}</em><div>{_esc(x.get('degree'))}</div></div>" for x in cv.get("education", []))
    projects = "".join(f"<div class='item'><b>{_esc(x.get('name'))}</b> <span>{_esc(x.get('link'))}</span><ul>{''.join(f'<li>{_esc(b)}</li>' for b in x.get('bullets', []))}</ul></div>" for x in cv.get("projects", []))
    links = " · ".join(f"{_esc(x.get('label'))}: {_esc(x.get('url'))}" for x in cv.get("links", []))
    return f"""<style>.cv{{font-family:Arial,sans-serif;color:#172033;line-height:1.45;padding:26px 34px;background:white}}.cv h1{{font-size:29px;margin:0;color:#101827}}.cv .headline{{color:{accent};font-size:15px;margin:4px 0 8px}}.cv .contact{{font-size:11px;color:#4b5563}}.cv h2{{font-size:12px;text-transform:uppercase;letter-spacing:.8px;color:{accent};border-bottom:1px solid #dbe0ea;padding-bottom:4px;margin:18px 0 7px}}.cv p{{font-size:12px;margin:0}}.item{{font-size:12px;margin:8px 0}}.item span{{color:#4b5563}}.item em{{float:right;color:#4b5563;font-size:11px;font-style:normal}}.item ul{{margin:3px 0 0;padding-left:18px}}.item li{{margin:2px 0}}.skills{{display:flex;flex-wrap:wrap;gap:5px}}.skill{{font-size:11px;background:#eef2ff;color:{accent};padding:3px 7px;border-radius:10px}}</style><div class='cv'><h1>{_esc(cv.get('name'))}</h1><div class='headline'>{_esc(cv.get('headline'))}</div><div class='contact'>{' · '.join(_esc(x) for x in cv.get('contact', []))}</div><div class='contact'>{links}</div>{section('Professional Summary', f'<p>{_esc(cv.get("summary"))}</p>')}{section('Experience', experience)}{section('Education', education)}{section('Skills', '<div class="skills">' + ''.join(f'<span class="skill">{_esc(x)}</span>' for x in cv.get('skills', [])) + '</div>')}{section('Projects', projects)}</div>"""


def export_docx(cv: dict[str, Any]) -> bytes:
    doc = Document(); sec = doc.sections[0]; sec.top_margin = sec.bottom_margin = Inches(.55)
    style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(10)
    doc.add_heading(cv.get('name', ''), 0); doc.add_paragraph(cv.get('headline', ''))
    doc.add_paragraph(' | '.join(cv.get('contact', [])))
    def add(title, items):
        if not items: return
        doc.add_heading(title, level=1)
        for item in items: doc.add_paragraph(item)
    add('Professional Summary', [cv.get('summary', '')])
    for title, entries, fields in [('Experience', cv.get('experience', []), ('title','company','dates','bullets')), ('Education', cv.get('education', []), ('school','degree','dates',None)), ('Projects', cv.get('projects', []), ('name','link',None,'bullets'))]:
        if entries:
            doc.add_heading(title, level=1)
            for x in entries:
                doc.add_paragraph(' | '.join(str(x.get(k,'')) for k in fields[:3] if k), style='Heading 2')
                for bullet in x.get('bullets', []): doc.add_paragraph(bullet, style='List Bullet')
    add('Skills', [', '.join(cv.get('skills', []))]); add('Links', [f"{x.get('label')}: {x.get('url')}" for x in cv.get('links', [])])
    output = io.BytesIO(); doc.save(output); return output.getvalue()


def _has_arabic(value: Any) -> bool:
    return bool(re.search(r"[\u0600-\u06FF]", str(value)))


def _pdf_text(value: Any, rtl: bool) -> str:
    text = _esc(value)
    if rtl:
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
            return get_display(arabic_reshaper.reshape(text))
        except Exception:
            return text
    return text


def export_pdf(cv: dict[str, Any]) -> bytes:
    output = io.BytesIO(); doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=1.5*cm, leftMargin=1.5*cm, topMargin=1.25*cm, bottomMargin=1.25*cm)
    all_text = json.dumps(cv, ensure_ascii=False)
    rtl = _has_arabic(all_text)
    regular, bold = 'Helvetica', 'Helvetica-Bold'
    font_path = Path('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf')
    bold_path = Path('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf')
    if rtl and font_path.exists():
        pdfmetrics.registerFont(TTFont('CVArabic', str(font_path)))
        pdfmetrics.registerFont(TTFont('CVArabicBold', str(bold_path if bold_path.exists() else font_path)))
        regular, bold = 'CVArabic', 'CVArabicBold'
    styles = getSampleStyleSheet(); accent = colors.HexColor('#4f46e5')
    alignment = TA_LEFT
    name = ParagraphStyle('name', parent=styles['Title'], fontName=bold, fontSize=22, leading=25, textColor=colors.HexColor('#111827'), alignment=alignment)
    heading = ParagraphStyle('heading', parent=styles['Heading2'], fontName=bold, fontSize=11, textColor=accent, spaceBefore=11, spaceAfter=4, alignment=alignment)
    body = ParagraphStyle('body', parent=styles['BodyText'], fontName=regular, fontSize=9.5, leading=13, alignment=alignment)
    story = [Paragraph(_pdf_text(cv.get('name'), rtl), name), Paragraph(_pdf_text(cv.get('headline'), rtl), ParagraphStyle('role', parent=body, textColor=accent)), Paragraph(_pdf_text(' | '.join(cv.get('contact', [])), rtl), body), Spacer(1, 5)]
    def title(t): story.extend([Paragraph(t, heading), HRFlowable(width='100%', thickness=.5, color=colors.HexColor('#d1d5db')), Spacer(1, 3)])
    if cv.get('summary'): title('PROFESSIONAL SUMMARY'); story.append(Paragraph(_pdf_text(cv['summary'], rtl), body))
    if cv.get('experience'):
        title('EXPERIENCE')
        for x in cv['experience']:
            story.append(Paragraph(f"<b>{_pdf_text(x.get('title'), rtl)}</b> · {_pdf_text(x.get('company'), rtl)} — {_pdf_text(x.get('dates'), rtl)}", body))
            for b in x.get('bullets', []): story.append(Paragraph('• ' + _pdf_text(b, rtl), body))
    if cv.get('education'):
        title('EDUCATION')
        for x in cv['education']: story.append(Paragraph(f"<b>{_pdf_text(x.get('school'), rtl)}</b> — {_pdf_text(x.get('degree'), rtl)} ({_pdf_text(x.get('dates'), rtl)})", body))
    if cv.get('skills'): title('SKILLS'); story.append(Paragraph(_pdf_text(' · '.join(cv['skills']), rtl), body))
    if cv.get('projects'):
        title('PROJECTS')
        for x in cv['projects']:
            story.append(Paragraph(f"<b>{_pdf_text(x.get('name'), rtl)}</b> · {_pdf_text(x.get('link'), rtl)}", body))
            for b in x.get('bullets', []): story.append(Paragraph('• ' + _pdf_text(b, rtl), body))
    doc.build(story); return output.getvalue()
